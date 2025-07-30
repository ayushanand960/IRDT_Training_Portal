import os
import time
import pandas as pd
from tqdm import tqdm
from docx import Document
from datetime import datetime
from django.conf import settings
from django.core.files.storage import default_storage
from Training.models import TrainingProgram
from Certificate.models import Certificate
from Login.models import User
from Certificate.utils.utils import replace_placeholders
import pythoncom
import win32com.client
from zipfile import ZipFile


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert a DOCX file to PDF using Microsoft Word COM automation.
    Retries 3 times in case Word crashes or fails to initialize.
    """
    try:
        pythoncom.CoInitialize()
        for attempt in range(3):
            try:
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
                doc.Close()
                word.Quit()
                # Check PDF was successfully created and has size > 1 KB
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1000:
                    return True
            except Exception as e:
                print(f"⚠️ Word attempt {attempt + 1} failed: {e}")
                time.sleep(1)
        print(f"❌ PDF generation failed for {docx_path}")
        return False
    except Exception as e:
        print(f"❌ Error initializing Word: {e}")
        return False


def generate_certificates_from_excel(file_path, template_path, training_code, coordinator_user):
    """
    Main function to generate certificates in PDF format from an Excel file and Word template.

    Steps:
    - Read user data from Excel
    - Match user via ehrms_code or email
    - Replace placeholders in Word template
    - Convert to PDF
    - Remove any temp/old certificates
    - Save the final clean PDF to disk and database
    """
    # Initialize COM if not already done
    try:
        pythoncom.CoInitialize()
    except:
        pass

    # Load trainee data from Excel
    try:
        df = pd.read_excel(file_path)
    except Exception as e:
        print(f"❌ Error reading Excel: {e}")
        return []

    # Get training program
    try:
        training = TrainingProgram.objects.get(code=training_code)
    except TrainingProgram.DoesNotExist:
        print(f"❌ Training with code '{training_code}' not found.")
        return []

    # Check if Word template exists
    if not os.path.exists(template_path):
        print(f"❌ Template not found at: {template_path}")
        return []

    # Directory to save final PDFs
    output_dir = os.path.join(settings.MEDIA_ROOT, 'certificates')
    os.makedirs(output_dir, exist_ok=True)

    generated_certificates = []

    # Iterate over each row in Excel
    for index, row in tqdm(df.iterrows(), total=len(df)):

        ehrms_code = row.get('ehrms_code') or row.get('EHRMS')
        email = row.get('email') or row.get('Email')

        # Match user using ehrms_code or email
        user = None
        if ehrms_code:
            user = User.objects.filter(ehrms_code=ehrms_code).first()
        if not user and email:
            user = User.objects.filter(email=email).first()

        if not user:
            print(f"⚠️ Skipping row — no matching user for EHRMS: {ehrms_code}, Email: {email}")
            continue

        try:
            # Load Word certificate template
            doc = Document(template_path)

            # Extract user and training details
            name = f"{user.first_name} {user.middle_name or ''} {user.last_name}".strip()
            designation = user.designation or ''
            branch = user.branch or ''
            institution = user.institute_name or ''
            year = training.start_date.year if training.start_date else datetime.now().year
            start_date = training.start_date or datetime.now()
            day = start_date.strftime("%d")         # 14
            month = start_date.strftime("%m")       # 06
            year_suffix = start_date.strftime("%y") # 25
            serial = str(index + 1).zfill(2)        # 2-digit serial number like 01, 10, etc.
            reference_number = f"{training.code}-{day}{month}{year_suffix}-{serial}"

            # Define replacements for placeholders
            replacements = {
                '{{name of staff}}': name,
                '{{designation}}': designation,
                '{{branch}}': branch,
                '{{institute name}}': institution,
                '{{certificate no}}': reference_number,
            }

            # Replace placeholders in paragraphs and tables
            for p in doc.paragraphs:
                replace_placeholders(p, replacements)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders(p, replacements)

            # Define clean file names (no temp suffixes)
            docx_name = f"{user.ehrms_code}_{training.code}_{year}.docx"
            pdf_name = docx_name.replace('.docx', '.pdf')
            db_file_name = pdf_name
            storage_path = f"certificates/{db_file_name}"
            docx_path = os.path.join(output_dir, docx_name)
            pdf_path = os.path.join(output_dir, pdf_name)

            # Delete existing certificate entries and files if they exist
            existing_certs = Certificate.objects.filter(user=user, training=training)
            for cert in existing_certs:
                if cert.certificate_file and os.path.exists(cert.certificate_file.path):
                    try:
                        os.remove(cert.certificate_file.path)
                    except Exception as e:
                        print(f"⚠️ Failed to delete previous file: {e}")
                cert.delete()

            # Remove existing file in storage if already exists
            if default_storage.exists(storage_path):
                default_storage.delete(storage_path)

            # Save updated .docx temporarily
            doc.save(docx_path)

            print(f"📄 Converting: {pdf_path}")
            success = convert_docx_to_pdf(docx_path, pdf_path)

            # Clean up the temporary DOCX file
            if os.path.exists(docx_path):
                os.remove(docx_path)

            if not success:
                print(f"❌ Skipping {user.email}: PDF conversion failed.")
                continue

            # Wait briefly to ensure PDF is completely written
            for _ in range(10):
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1000:
                    break
                time.sleep(0.5)
            else:
                print(f"❌ Skipping {user.email}: PDF not found after wait.")
                continue

            # Read PDF content safely
            with open(pdf_path, 'rb') as f:
                file_bytes = f.read()

            # Delete physical PDF to avoid suffixes like 'temp0' etc.
            os.remove(pdf_path)

            # Create and save certificate object in DB
            cert = Certificate.objects.create(
            user=user,
            training=training,
            uploaded_by=coordinator_user,
            full_name=name,
            designation=designation,
            institution=institution,
            reference_number=reference_number,# ✅ Use the correct variable
            )


            # Save final cleaned PDF
            final_cert_path = os.path.join(settings.MEDIA_ROOT, 'certificates', db_file_name)
            with open(final_cert_path, 'wb') as f:
                f.write(file_bytes)

            # Update DB with correct file path
            cert.certificate_file.name = f"certificates/{db_file_name}"
            cert.save()

            generated_certificates.append(cert)

        except Exception as e:
            print(f"❌ Error for user {user.email or ehrms_code}: {e}")
            continue

    # Finalize COM objects
    try:
        pythoncom.CoUninitialize()
    except:
        pass

    zip_file_path = create_zip_for_training(training_code)

    return generated_certificates, zip_file_path

def create_zip_for_training(training_code):
    cert_folder = os.path.join(settings.MEDIA_ROOT, "certificates")
    zip_folder = os.path.join(settings.MEDIA_ROOT, "certificates", "zips")
    os.makedirs(zip_folder, exist_ok=True)

    zip_file_path = os.path.join(zip_folder, f"{training_code}_certificates.zip")

    with ZipFile(zip_file_path, 'w') as zipf:
        for filename in os.listdir(cert_folder):
            if filename.startswith(training_code) and filename.endswith('.pdf'):
                filepath = os.path.join(cert_folder, filename)
                zipf.write(filepath, arcname=filename)

    return zip_file_path




# import os
# import time
# import pandas as pd
# from tqdm import tqdm
# from docx import Document
# from datetime import datetime
# from django.conf import settings
# from django.core.files.storage import default_storage
# from Training.models import TrainingProgram
# from Certificate.models import Certificate
# from Login.models import User
# from Certificate.utils.utils import replace_placeholders
# import pythoncom
# import win32com.client


# def convert_docx_to_pdf(docx_path, pdf_path):
#     try:
#         pythoncom.CoInitialize()
#         word = win32com.client.DispatchEx("Word.Application")
#         word.Visible = False
#         doc = word.Documents.Open(docx_path)
#         doc.SaveAs(pdf_path, FileFormat=17)
#         doc.Close()
#         word.Quit()

#         if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1000:
#             return True
#     except Exception as e:
#         print(f"❌ Word conversion failed: {e}")
#     return False


# def get_training_short_name(name):
#     return ''.join(word[0].upper() for word in name.split() if word.isalpha())


# def get_certificate_reference(training, serial_no):
#     try:
#         start_date = training.start_date
#         day = start_date.day
#         month_initial = start_date.strftime('%B')[0].upper()
#         year_suffix = str(start_date.year)[-2:]
#         short_name = get_training_short_name(training.name)
#         return f"{short_name}-{day}{month_initial}{year_suffix}-{serial_no}"
#     except Exception as e:
#         print(f"⚠️ Failed to generate reference number: {e}")
#         return "UNKNOWN-REF"


# def generate_certificates_from_excel(file_path, template_path, training_code, coordinator_user):
#     try:
#         pythoncom.CoInitialize()
#     except:
#         pass

#     try:
#         df = pd.read_excel(file_path)
#     except Exception as e:
#         print(f"❌ Error reading Excel: {e}")
#         return []

#     try:
#         training = TrainingProgram.objects.get(code=training_code)
#     except TrainingProgram.DoesNotExist:
#         print(f"❌ Training with code '{training_code}' not found.")
#         return []

#     if not os.path.exists(template_path):
#         print(f"❌ Template not found at: {template_path}")
#         return []

#     output_dir = os.path.join(settings.MEDIA_ROOT, 'certificates')
#     os.makedirs(output_dir, exist_ok=True)

#     generated_certificates = []

#     for index, row in tqdm(df.iterrows(), total=len(df)):
#         ehrms_code = row.get('ehrms_code') or row.get('EHRMS')
#         email = row.get('email') or row.get('Email')

#         user = None
#         if ehrms_code:
#             user = User.objects.filter(ehrms_code=str(ehrms_code).strip()).first()
#         if not user and email:
#             user = User.objects.filter(email=str(email).strip()).first()

#         if not user:
#             print(f"⚠️ Skipping row {index + 1} — no matching user for EHRMS: {ehrms_code}, Email: {email}")
#             continue

#         try:
#             doc = Document(template_path)

#             name = f"{user.first_name} {user.middle_name or ''} {user.last_name}".strip()
#             designation = user.designation or ''
#             branch = user.branch or ''
#             institution = user.institute_name or ''
#             year = training.start_date.year if training.start_date else datetime.now().year

#             serial_no = index + 1
#             reference_number = get_certificate_reference(training, serial_no)

#             replacements = {
#                 '{{name of staff}}': name,
#                 '{{designation}}': designation,
#                 '{{branch}}': branch,
#                 '{{institute name}}': institution,
#                 '{{certificate no}}': reference_number,
#             }

#             for p in doc.paragraphs:
#                 replace_placeholders(p, replacements)
#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         for p in cell.paragraphs:
#                             replace_placeholders(p, replacements)

#             docx_name = f"{user.ehrms_code}_{training.code}_{year}.docx"
#             pdf_name = docx_name.replace('.docx', '.pdf')
#             final_pdf_path = os.path.join(output_dir, pdf_name)

#             # Delete old certificates
#             old_certs = Certificate.objects.filter(user=user, training=training)
#             for cert in old_certs:
#                 try:
#                     if cert.certificate_file and os.path.exists(cert.certificate_file.path):
#                         os.remove(cert.certificate_file.path)
#                     cert.delete()
#                 except Exception as e:
#                     print(f"⚠️ Failed to delete old cert: {e}")

#             doc.save(os.path.join(output_dir, docx_name))
#             success = convert_docx_to_pdf(os.path.join(output_dir, docx_name), final_pdf_path)
#             os.remove(os.path.join(output_dir, docx_name))  # clean .docx

#             if not success:
#                 print(f"❌ PDF generation failed for {user.email}")
#                 continue

#             # Save to DB
#             with open(final_pdf_path, 'rb') as f:
#                 cert = Certificate.objects.create(
#                     user=user,
#                     training=training,
#                     uploaded_by=coordinator_user,
#                     full_name=name,
#                     designation=designation,
#                     institution=institution,
#                     reference_number=reference_number,
#                 )
#                 cert.certificate_file.save(pdf_name, f)
#                 cert.save()

#             os.remove(final_pdf_path)  # Clean temp PDF

#             generated_certificates.append(cert)

#         except Exception as e:
#             print(f"❌ Error for user {user.email or ehrms_code}: {e}")
#             continue

#     try:
#         pythoncom.CoUninitialize()
#     except:
#         pass

#     return generated_certificates








